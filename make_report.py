#!/usr/bin/env python3
"""Generate the findings summary as a Word document (Japanese).

v4 (2026-07-09): trip-extraction corruption fixed (notes/review_2026-07-09.md),
all data regenerated (54/5,486 winners vs 156/31,906 losers), control-group
revision folded in. Numbers in v3 and earlier are superseded.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/home/smoltz/claude/crypto_whales/調査まとめ_2026-07-09.docx"

doc = Document()

# --- base styles -----------------------------------------------------------
style = doc.styles["Normal"]
style.font.name = "Noto Sans CJK JP"
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(
    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
    "Noto Sans CJK JP")

for lvl, sz in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)):
    h = doc.styles[lvl]
    h.font.name = "Noto Sans CJK JP"
    h.font.size = Pt(sz)
    h.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    h.element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
        "Noto Sans CJK JP")


def p(text, bold=False, size=None, style_name=None):
    para = doc.add_paragraph(style=style_name)
    run = para.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    return para


def bullet(text, bold_prefix=None):
    para = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = para.add_run(bold_prefix)
        r.bold = True
        para.add_run(text)
    else:
        para.add_run(text)
    return para


def numbered(text, bold_prefix=None):
    para = doc.add_paragraph(style="List Number")
    if bold_prefix:
        r = para.add_run(bold_prefix)
        r.bold = True
        para.add_run(text)
    else:
        para.add_run(text)
    return para


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for i, row in enumerate(rows, 1):
        for j, v in enumerate(row):
            t.rows[i].cells[j].text = str(v)
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Cm(w)
    return t


# --- title -------------------------------------------------------------------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Hyperliquid スキルトレーダー & BTC オンチェーン調査\nまとめと提言（v4）")
r.bold = True
r.font.size = Pt(18)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("2026年7月9日（抽出バグ修正・全データ再生成・対照群反映版）"
                 " / crypto_whales プロジェクト")
sr.font.size = Pt(10)
sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

note = doc.add_paragraph()
nr = note.add_run("本版で v3 以前の数値は全て無効です。7/9 のコードレビューでトリップ抽出の"
                  "データ破損バグ（取引履歴の欠落区間を跨いだトレードの誤結合、243口座中87で発生）を"
                  "発見・修正し、全データを再生成・全検定をやり直しました。結論も一部変わっています。")
nr.font.size = Pt(9.5)
nr.font.color.rgb = RGBColor(0x8b, 0x00, 0x00)

# --- 1. overview ---------------------------------------------------------------
doc.add_heading("1. このプロジェクトは何か", level=1)
p("2つのデータ領域を並行して調べています。(1) Hyperliquid（無期限先物DEX）で"
  "「本当に上手いトレーダー」を口座規模に関係なく発見し、その手法を分析する"
  "（コピートレードではなく手法研究）。(2) Bitcoin ベースチェーン上で取引所など"
  "既知エンティティの大口資金移動を監視する。どちらも無料の公開APIのみで運用中です。")

# --- 2. findings ---------------------------------------------------------------
doc.add_heading("2. わかったこと", level=1)

doc.add_heading("2.1 スキルは実在し、持続する【確度: 高】", level=2)
p("リーダーボードの損益ランキングは当てにならない一方、約定履歴（fills）から計算した"
  "実測成績は将来の成績をよく予測します。125口座の分割検証"
  "（7/9 に手数料計上バグを修正して再検定、結論は不変）では：")
bullet("fills 由来の指標（勝率・純損益・プロフィットファクター）と後半成績の相関は "
       "+0.59〜+0.64。リーダーボード損益は +0.16 しかない。")
bullet("前半勝率が上位半分の口座は 82% が後半もプラス（下位半分は 30%、"
       "当てずっぽうの基準は 56%）。")
p("→「whale = 資金量」ではなく「whale = スキル」で定義すべき、という当初仮説を確認。", bold=True)

doc.add_heading("2.2 勝者集団の行動（BTC・5,486トリップ / 54人）【確度: 高】", level=2)
p("検証済みスキル保有者のうち、建玉をフラットに戻す「方向性トレーダー」全員の"
  "個別トレード（トリップ）を全取引履歴（1人あたり最大30,000件をページング取得・"
  "欠落区間を跨ぐトリップは除外）から抽出し、セットアップを再検定した結果：")

table(
    ["観察", "内容", "統計検定の結果"],
    [
        ["① 損切りは速く、勝ちは伸ばす",
         "負けトレードの保有時間は勝ちの0.35倍（中央値）。全体では負け42分 vs 勝ち"
         "2.1時間。勝率は47%しかないのに非対称なペイオフで儲けている。",
         "勝者集団内では確立（42/51人、合成 p≈7×10⁻³⁸、95%CI [0.22, 0.54]）。"
         "ただし §2.3 のとおり敗者も同じことをする — 勝者の見分け方にはならない"],
        ["② 追いかけず、逆張りする",
         "直前4時間の値動きと同方向のエントリーは39%のみ。ロングは中央値−0.36%の"
         "下落後に入る。『逆張りの方が順張りより勝ちやすい』かは、破損データでは"
         "不成立だったが、クリーンデータで名目的に有意になった。",
         "行動としては確か。エッジとしては弱い証拠が再浮上"
         "（CMH OR 1.19, p=0.023 — 多重比較・データ間で判定が反転した経緯から、"
         "フォワード検証を通るまで確定扱いしない）"],
        ["③ ボラティリティが引き金",
         "エントリーの38%は4時間で1%以上動いた直後（勝率 52% vs 静穏時 49%）。",
         "不成立のまま（OR 1.05, p=0.53）"],
        ["④ 積み増し（ピラミッディング）は常態",
         "73%のトリップでポジションに追加（中央値8回）。『追加した方が勝ちやすい』"
         "というトリップ単位のエッジはない。",
         "行動としては確か。トリップ単位のエッジは不成立（OR 0.97, p=0.71）。"
         "ただし §2.3 のとおりトレーダー単位では勝者ほど積み増す"],
        ["⑤ 時間帯はNYセッション",
         "エントリーはUTC 13〜17時（日本時間22〜2時、NY市場の午前）に強く集中"
         "（ピークは14〜15時UTC）。",
         "記述統計のみ"],
    ],
    widths=[4.0, 8.0, 5.0],
)
p("")
p("統計処理の要点：プール集計は高頻度口座に引きずられるため、勝率の比較はすべて"
  "トレーダー層別の Cochran-Mantel-Haenszel 検定で行いました。破損データで一度"
  "「死んだ」逆張りエッジがクリーンデータで蘇った事実は、この種の判定が抽出品質に"
  "敏感であることの実例です — だからこそ②はフォワード検証待ちとしています。")

doc.add_heading("2.3 対照群との比較 — 本調査の最重要結論【確度: 高】", level=2)
p("同じ活動量で成績が逆（純損失）の検証済み敗者 156人 / 31,906トリップに同じ分析をかけ、"
  "勝者 54人と比較しました（トレーダー単位の Mann-Whitney U、両側）：")
table(
    ["特徴", "勝者", "敗者", "p値", "判定"],
    [
        ["損切り実行率（負け保有<勝ち保有の人の割合）", "82%", "82%", "—",
         "同率 — 全員やっている"],
        ["損切りの程度（負け/勝ち保有比の中央値）", "0.35", "0.47", "0.29",
         "有意差なし"],
        ["勝ち保有時間の中央値", "2.8h", "1.4h", "0.056", "傾向どまり"],
        ["トレード数（中央値）", "48.5", "99", "0.001", "頑健 — 敗者は2倍取引する"],
        ["積み増し率", "89%", "81%", "0.004", "頑健"],
        ["建玉サイズ（中央値）", "$115K", "$49K", "0.006", "頑健 — 2.3倍"],
        ["逆張り率", "56%", "53%", "0.060", "傾向どまり"],
    ],
    widths=[6.0, 2.2, 2.2, 1.8, 4.8],
)
p("")
p("損切り規律は勝者の特徴ではなく、この市場で生き残っている全員の共通動作でした"
  "（旧版の「勝者は損切りが速い」「勝者は勝ちを切らない」はどちらも降格）。"
  "クリーンデータで頑健に残った勝者の見分け方は：", bold=False)
p("少なく張る（半分の頻度）・大きく張る（2.3倍）・積み増して張る。", bold=True)
p("裏返すと、敗者の最も確かなマーカーは過剰トレード（2倍の頻度で半分のサイズ）です。"
  "なお勝率・ペイオフ比の群間差は「勝者/敗者を損益で選抜した」ことの裏返しを含むため、"
  "発見としては引用しません。")

doc.add_heading("2.4 「集団の方向」にはエッジがない【確度: 中】", level=2)
p("スキル加重した whale 集団の BTC/ETH ポジション方向は、1〜24時間先のリターンを"
  "予測しません（全ホライズンでベースレート以下）。上手い個人の手法を学ぶのが正しく、"
  "集団の売買方向を真似るのは（現時点では）無意味 — プロジェクトの「コピートレード"
  "ではない」方針を裏付ける結果です。30日分たまる7月末に再検証します"
  "（7/9 に劣化スナップショット除外の頑健化も実施済み）。")

doc.add_heading("2.5 BTC オンチェーン監視【確度: 高（インフラ）】", level=2)
bullet("チェーンの可視性は無料・完全。難しいのは「どのアドレスが誰か」（帰属）だけ。")
bullet("無料の GraphSense ラベルで大手取引所はカバー可能。監視対象は 11エンティティ / "
       "約76万BTC（Binance・Bitfinex・OKX・Crypto.com・Huobi・Bybit のコールド/ホット、"
       "全て残高を実測検証済み、7/5 拡張）。毎時 cron で時系列を蓄積中。")
bullet("教訓:「ラベルは腐る、残高が真実」。ラベル上は取引所ホットウォレットでも残高"
       "ゼロが多数（ローテーション済み）。残高検証なしにアドレスを信用しない。")
bullet("BlackRock / Coinbase / MicroStrategy の帰属だけは無料データでは不可能 — "
       "Arkham の無料APIキーが必要（コードは実装済み・キー待ち）。")

# --- 3. useful ---------------------------------------------------------------
doc.add_heading("3. 有益なこと（実用的な示唆）", level=1)
numbered("速い損切りは勝者も敗者も同率（82%）で実践しており、それだけでは勝てません。"
         "ただし『やらない』選択肢はない — 規律は利益の源泉ではなく参加条件です。",
         bold_prefix="損切り規律は前提条件であって差別化要因ではない。")
numbered("クリーンデータで勝者と敗者を頑健に分けたのは行動の型ではなく選択性とサイズ："
         "取引頻度が半分（p=0.001）、建玉2.3倍（p=0.006）、積み増し率が高い（p=0.004）。"
         "個人のトレードに移植するなら「トレードを半分に間引き、確信のある玉に集中して"
         "分割で積む」が本調査の中心的な示唆です。",
         bold_prefix="少なく・大きく・積み増して張る。")
numbered("「勝ちを伸ばす」（勝ち保有 2.8h vs 1.4h, p=0.056）と「逆張りエッジ」"
         "（p=0.023、判定が一度反転）は有望ですが、フォワード検証（7月下旬）を通るまで"
         "ルール化しないでください。",
         bold_prefix="傾向どまりの2項目は確定扱いしない。")
numbered("集団の売買方向をシグナルとして使うのは現時点で根拠なし。オンチェーンでは"
         "ホットウォレットのネットフローだけが売買圧のシグナルで、コールドの動きは"
         "保管の付け替え（ノイズ）です。",
         bold_prefix="群衆シグナルは使わない。見るのはホットウォレットのみ。")

# --- 4. actions ---------------------------------------------------------------
doc.add_heading("4. あなたが行うべきこと", level=1)

doc.add_heading("4.1 すぐできる（所要10分・無料）", level=2)
numbered("intel.arkm.com/api で無料アカウントを作り、API キーを取得して環境変数 "
         "ARKHAM_API_KEY に設定する。これだけで BlackRock→Coinbase 級の資金フロー追跡"
         "（arkham_flows.py）が即座に動きます。現在唯一、人手が必要な項目です。",
         bold_prefix="Arkham 無料キーの取得。")

doc.add_heading("4.2 待つだけ（cron が自動で蓄積中）", level=2)
table(
    ["いつ", "何を", "目的"],
    [
        ["2026-07-19〜08-02", "skilled.py 再実行（フォワード検証）",
         "検証済みトレーダーの生存確認 + §2.3 の「少なく・大きく」仮説と"
         "傾向どまり2項目の再検定"],
        ["2026-07-27 頃", "whales_track.py --analyze --horizon 1,4,8,12,24",
         "集団方向シグナルの30日再検証（現状エッジなしの確認）"],
        ["2026-08-05 頃", "btc_track.py --analyze --horizon 1,4,24",
         "取引所ネットフロー→BTCリターン予測の初検証"],
    ],
    widths=[3.5, 7.5, 6.0],
)

doc.add_heading("4.3 判断してほしいこと", level=2)
numbered("CLI・アラート・ダッシュボードのどれにするかは未決のまま（意図的）。"
         "フォワード検証と30日再検証の結果が出る8月頭に決めるのが合理的です。",
         bold_prefix="出力形式はまだ決めなくてよい。")

# --- 5. caveats ---------------------------------------------------------------
doc.add_heading("5. 注意（読み違えないために）", level=1)
bullet("データ品質の経緯: 7/9 のレビューで、取引履歴の欠落区間（ページング上限・"
       "約定を伴わない建玉変化）を跨いだトリップが誤結合されるバグを発見しました"
       "（最悪例は97日間を1トレード扱い）。修正後の抽出では欠落区間を跨ぐトリップを"
       "除外し、上限を30,000件に拡大して全数値を再導出しています。本版以前の"
       "レポートの数値は使わないでください。")
bullet("勝率・ペイオフ比の勝者/敗者差は選抜の裏返しを含む（群を損益で定義したため）。"
       "引用してよいのは行動系の差（頻度・サイズ・積み増し・保有時間）だけです。")
bullet("トリップは同一トレーダー内で独立ではない（時間的クラスタ・同一相場）ため、"
       "p値はやや楽観的です。クロストレーダーの主張の実効サンプルは勝者54人・敗者156人。")
bullet("②の逆張りエッジは3つの勝率コントラストの1つ（多重比較）で、Bonferroni 補正"
       "後は p≈0.07。データ再生成で判定が反転した経緯も含め、弱い証拠として扱うこと。")
bullet("勝率のインフレ: 含み損を抱えたまま実現益だけ計上する口座は勝率が過大に出ます。"
       "プロフィットファクターと純損益を優先して見ています。")
bullet("フェッチ失敗 101/836 口座分は今回の母集団に入っていません（ウォームキャッシュ"
       "再実行で回収可能 — 母集団がわずかに増える可能性があります）。")

doc.save(OUT)
print(f"saved: {OUT}")
